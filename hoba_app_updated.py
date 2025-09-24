
# JIBC Workshops Navigator (updated v6)
# - Adds: image zoom slider (1.0x–3.0x) and viewer height slider (500–1400px)
# - PDF display toggle: Auto (Images → Data-URI → File URL), Images only, Data-URI only, File URL only
# - Left nav for DISCOVERY / DESIGN / REQUIREMENTS / INTEGRATION / ROADMAP / FINALIZATION
# - Robust DOCX parsing (bytes or path) with forgiving phase heading detection

import os
import io
import re
import base64
import tempfile
from typing import Dict, List, Optional

import streamlit as st
from streamlit.components.v1 import html

# Optional imports for PDF->image and DOCX parsing
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

DEFAULT_PDF_PATH = "/mnt/data/JIBC_Workshops_Facilitator_Dashboard_Phased_Timeline.pdf"
DEFAULT_DOCX_PATH = "/mnt/data/JIBC_Workshop_Facilitation_Guidebook_Full.docx"

SECTIONS = ["DISCOVERY", "DESIGN", "REQUIREMENTS", "INTEGRATION", "ROADMAP", "FINALIZATION"]

PHASE_PATTERNS = {
    "DISCOVERY": re.compile(r"\bdiscovery\b", re.I),
    "DESIGN": re.compile(r"\bdesign\b", re.I),
    "REQUIREMENTS": re.compile(r"\brequirement(s)?\b", re.I),
    "INTEGRATION": re.compile(r"\bintegration(s)?\b", re.I),
    "ROADMAP": re.compile(r"\broadmap\b", re.I),
    "FINALIZATION": re.compile(r"\bfinali[sz]ation|final readout|approval\b", re.I),
}


def pdf_to_images(pdf_path: Optional[str] = None, pdf_bytes: Optional[bytes] = None, zoom: float = 2.0) -> List[bytes]:
    """Convert a PDF (from path or bytes) into a list of PNG images (bytes)."""
    if fitz is None:
        return []
    try:
        if pdf_bytes is not None:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        elif pdf_path and os.path.exists(pdf_path):
            doc = fitz.open(pdf_path)
        else:
            return []
        images = []
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("png"))
        return images
    except Exception as e:
        st.error(f"PDF render error: {e}")
        return []


def embed_pdf_as_data_uri(pdf_path: Optional[str] = None, pdf_bytes: Optional[bytes] = None, height: int = 820):
    """Embed a PDF using a base64 data URI so the browser can display it without a file URL."""
    try:
        if pdf_bytes is None:
            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()
            else:
                st.warning("No PDF available to display.")
                return
        b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        data_uri = f"data:application/pdf;base64,{b64}"
        html_code = f"""
        <div style='border:1px solid #ddd;border-radius:8px;overflow:hidden;'>
          <iframe src="{data_uri}" width="100%" height="{height}px" style="border:0;"></iframe>
        </div>
        """
        html(html_code, height=height + 10)
    except Exception as e:
        st.error(f"PDF embed error (data-URI): {e}")


def embed_pdf_as_file_url(pdf_path: Optional[str] = None, pdf_bytes: Optional[bytes] = None, height: int = 820):
    """Embed a PDF using a file URL (or a temp file if only bytes are provided)."""
    try:
        if pdf_bytes is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tf:
                tf.write(pdf_bytes)
                path_to_show = tf.name
        elif pdf_path and os.path.exists(pdf_path):
            path_to_show = pdf_path
        else:
            st.warning("No PDF available to display.")
            return

        html_code = f"""
        <div style='border:1px solid #ddd;border-radius:8px;overflow:hidden;'>
          <embed src="{path_to_show}" type="application/pdf" width="100%" height="{height}px" />
        </div>
        """
        html(html_code, height=height + 10)
    except Exception as e:
        st.error(f"PDF embed error (file URL): {e}")


def extract_sections_from_docx(docx_path: Optional[str] = None, docx_bytes: Optional[bytes] = None) -> Dict[str, str]:
    """Parse the Word doc (from path or bytes) and extract text under six top-level sections."""
    out = {s: [] for s in SECTIONS}
    if DocxDocument is None:
        return {s: "*(python-docx not available — install with `pip install python-docx`)*" for s in SECTIONS}

    try:
        if docx_bytes is not None:
            doc = DocxDocument(io.BytesIO(docx_bytes))
        elif docx_path and os.path.exists(docx_path):
            doc = DocxDocument(docx_path)
        else:
            return {s: f"*(DOCX not found at {docx_path or '[uploaded bytes missing]'})*" for s in SECTIONS}
    except Exception as e:
        return {s: f"*(Failed to open DOCX: {e})*" for s in SECTIONS}

    current_section: Optional[str] = None

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        style_name = getattr(getattr(para, "style", None), "name", "") or ""

        # 1) Heading-based detection: any heading that matches a phase pattern
        if "HEADING" in style_name.upper():
            matched = False
            for sec, pattern in PHASE_PATTERNS.items():
                if pattern.search(text):
                    current_section = sec
                    out[sec].append(f"### {text}")
                    matched = True
                    break
            if matched:
                continue

        # 2) Non-heading "pseudo heading": short/uppercase lines with phase words
        for sec, pattern in PHASE_PATTERNS.items():
            if pattern.search(text) and (text.isupper() or len(text.split()) <= 6):
                current_section = sec
                out[sec].append(f"### {text}")
                break
        else:
            # 3) Regular content
            if current_section:
                out[current_section].append(text)

    return {k: ("\n\n".join(v) if v else "*(No section content found — check headings/keywords in the DOCX)*") for k, v in out.items()}


def show_sidebar_navigation():
    st.sidebar.title("Workshops Navigator")
    st.sidebar.caption("Click a phase to load content from the guidebook.")
    for sec in SECTIONS:
        if st.sidebar.button(sec, use_container_width=True, key=f"nav_{sec}"):
            st.session_state.current_section = sec
    st.sidebar.markdown("---")
    if st.sidebar.button("Show Dashboard PDF", use_container_width=True, key="show_pdf"):
        st.session_state.current_section = None


def diagnostics(pdf_path: Optional[str], docx_path: Optional[str], pdf_bytes: Optional[bytes], docx_bytes: Optional[bytes]):
    st.sidebar.markdown("### Diagnostics")
    st.sidebar.write(f"**python-docx available:** {'✅' if DocxDocument else '❌'}")
    st.sidebar.write(f"**PyMuPDF (fitz) available:** {'✅' if fitz else '❌'}")
    st.sidebar.write(f"**PDF path found:** {'✅' if (pdf_path and os.path.exists(pdf_path)) else '—'}")
    st.sidebar.write(f"**PDF uploaded:** {'✅' if pdf_bytes else '—'}")
    st.sidebar.write(f"**DOCX path found:** {'✅' if (docx_path and os.path.exists(docx_path)) else '—'}")
    st.sidebar.write(f"**DOCX uploaded:** {'✅' if docx_bytes else '—'}")


def show_pdf_view(pdf_path: Optional[str], pdf_bytes: Optional[bytes], mode: str, zoom: float, viewer_height: int):
    st.subheader("Facilitator’s Dashboard")
    st.caption("Switch PDF rendering modes and adjust zoom/height in the sidebar.")

    if mode == "Images only":
        images = pdf_to_images(pdf_path=pdf_path, pdf_bytes=pdf_bytes, zoom=zoom)
        if images:
            for i, img_bytes in enumerate(images, 1):
                st.image(img_bytes, caption=f"Dashboard Page {i}", use_container_width=True)
        else:
            st.warning("Could not render images (PyMuPDF missing or invalid PDF).")

    elif mode == "Data-URI only":
        embed_pdf_as_data_uri(pdf_path=pdf_path, pdf_bytes=pdf_bytes, height=viewer_height)

    elif mode == "File URL only":
        embed_pdf_as_file_url(pdf_path=pdf_path, pdf_bytes=pdf_bytes, height=viewer_height)

    else:  # Auto
        images = pdf_to_images(pdf_path=pdf_path, pdf_bytes=pdf_bytes, zoom=zoom)
        if images:
            for i, img_bytes in enumerate(images, 1):
                st.image(img_bytes, caption=f"Dashboard Page {i}", use_container_width=True)
        else:
            # Try Data-URI first, then File URL as last resort
            try:
                embed_pdf_as_data_uri(pdf_path=pdf_path, pdf_bytes=pdf_bytes, height=viewer_height)
            except Exception:
                embed_pdf_as_file_url(pdf_path=pdf_path, pdf_bytes=pdf_bytes, height=viewer_height)


def main():
    st.set_page_config(page_title="JIBC Workshops Navigator", page_icon="🗂️", layout="wide", initial_sidebar_state="expanded")
    if "current_section" not in st.session_state:
        st.session_state.current_section = None

    st.title("JIBC Workshops – Phased Navigator")
    st.write("View the phased dashboard PDF and jump to content sections from the full facilitation guidebook.")

    # Files expander: allow paths or uploads; prefer uploads when provided
    with st.sidebar.expander("Files", expanded=False):
        pdf_path = st.text_input("PDF path", value=DEFAULT_PDF_PATH)
        docx_path = st.text_input("DOCX path", value=DEFAULT_DOCX_PATH)
        uploaded_pdf = st.file_uploader("Or upload a PDF", type=["pdf"], key="pdf_up")
        uploaded_docx = st.file_uploader("Or upload a DOCX", type=["docx"], key="docx_up")

    pdf_bytes = uploaded_pdf.read() if uploaded_pdf is not None else None
    docx_bytes = uploaded_docx.read() if uploaded_docx is not None else None

    # Rendering controls
    mode = st.sidebar.radio(
        "PDF display mode",
        options=["Auto (Images → Data-URI → File URL)", "Images only", "Data-URI only", "File URL only"],
    )
    zoom = st.sidebar.slider("Image zoom (when rendering pages)", min_value=1.0, max_value=3.0, value=2.0, step=0.25)
    viewer_height = st.sidebar.slider("Viewer height (px) for Data-URI/File URL", min_value=500, max_value=1400, value=860, step=20)

    diagnostics(pdf_path, docx_path, pdf_bytes, docx_bytes)
    show_sidebar_navigation()

    # Extract sections from DOCX (from bytes if uploaded, else from path)
    sections_map = extract_sections_from_docx(docx_path if docx_bytes is None else None,
                                              docx_bytes=docx_bytes)

    # Show either the PDF (from bytes if uploaded) or the selected section
    if st.session_state.current_section is None:
        show_pdf_view(pdf_path if pdf_bytes is None else None, pdf_bytes=pdf_bytes, mode=mode, zoom=zoom, viewer_height=viewer_height)
    else:
        sec = st.session_state.current_section
        st.subheader(f"{sec} — Facilitation Guidebook Extract")
        content = sections_map.get(sec, "*(No content found for this section.)*")
        st.markdown(content)
        st.info("Use the sidebar to switch sections or click 'Show Dashboard PDF' to return to the PDF.")


if __name__ == "__main__":
    main()
