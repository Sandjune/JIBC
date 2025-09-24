
# JIBC Workshops Navigator (updated v9)
# - NEW: Display Word guidebook *as formatted*:
#     • Page images (best fidelity; requires LibreOffice/soffice or Word/docx2pdf)
#     • HTML via Mammoth (good for text/tables; shapes/boxes may vary)
#     • Extracted text (fastest; no formatting)
# - Existing: PDF display for dashboard (images/embeds); left-side phase navigation; zoom/height sliders
# - Auto-installs python-docx, pymupdf, mammoth (best effort)

import os
import io
import re
import base64
import tempfile
import importlib
import shutil
from typing import Dict, List, Optional

import streamlit as st
from streamlit.components.v1 import html

# Try imports (may be None initially)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# Optional Mammoth (DOCX -> HTML)
try:
    import mammoth
except Exception:
    mammoth = None

DEFAULT_PDF_PATH = "/mnt/data/JIBC_Workshops_Facilitator_Dashboard_Phased_Timeline.pdf"
DEFAULT_DOCX_PATH = "/mnt/data/JIBC_Workshop_Facilitation_Guidebook_Full.docx"

SECTIONS = ["DISCOVERY", "DESIGN", "REQUIREMENTS", "INTEGRATION", "ROADMAP", "FINALIZATION"]

PHASE_PATTERNS = {
    "DISCOVERY": re.compile(r"\bdiscovery\b", re.I),
    "DESIGN": re.compile(r"\bdesign\b", re.I),
    "REQUIREMENTS": re.compile(r"\brequirement(s)?\b", re.I),
    "INTEGRATION": re.compile(r"\bintegration(s)?\b", re.I),
    "ROADMAP": re.compile(r"\broadmap\b", re.I),
    "FINALIZATION": re.compile(r"\bfinali[sz]ation|final readout|approval\b", re.I),
}


def ensure_dependency(pip_name: str, import_name: str):
    """Ensure a dependency can be imported; pip install if needed (best effort)."""
    try:
        importlib.import_module(import_name)
        return True
    except Exception:
        try:
            import sys, subprocess
            with st.spinner(f"Installing {pip_name} ..."):
                subprocess.check_call([sys.executable, "-m", "pip", "install", pip_name])
            importlib.invalidate_caches()
            importlib.import_module(import_name)
            return True
        except Exception:
            return False


def pdf_to_images(pdf_path: Optional[str] = None, pdf_bytes: Optional[bytes] = None, zoom: float = 2.0) -> List[bytes]:
    """Convert a PDF (from path or bytes) into a list of PNG images (bytes)."""
    global fitz
    if fitz is None:
        try:
            fitz = importlib.import_module("fitz")
        except Exception:
            return []
    try:
        if pdf_bytes is not None:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        elif pdf_path and os.path.exists(pdf_path):
            doc = fitz.open(pdf_path)
        else:
            return []
        images = []
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            images.append(pix.tobytes("png"))
        return images
    except Exception as e:
        st.error(f"PDF render error: {e}")
        return []


def embed_pdf_as_data_uri(pdf_path: Optional[str] = None, pdf_bytes: Optional[bytes] = None, height: int = 820):
    """Embed a PDF using a base64 data URI so the browser can display it without a file URL."""
    try:
        if pdf_bytes is None:
            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()
            else:
                st.warning("No PDF available to display.")
                return
        b64 = base64.b64encode(pdf_bytes).decode("utf-8")
        data_uri = f"data:application/pdf;base64,{b64}"
        html_code = f"""
        <div style='border:1px solid #ddd;border-radius:8px;overflow:hidden;'>
          <iframe src="{data_uri}" width="100%" height="{height}px" style="border:0;"></iframe>
        </div>
        """
        html(html_code, height=height + 10)
    except Exception as e:
        st.error(f"PDF embed error (data-URI): {e}")


def embed_pdf_as_file_url(pdf_path: Optional[str] = None, pdf_bytes: Optional[bytes] = None, height: int = 820):
    """Embed a PDF using a file URL (or a temp file if only bytes are provided)."""
    try:
        if pdf_bytes is not None:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tf:
                tf.write(pdf_bytes)
                path_to_show = tf.name
        elif pdf_path and os.path.exists(pdf_path):
            path_to_show = pdf_path
        else:
            st.warning("No PDF available to display.")
            return

        html_code = f"""
        <div style='border:1px solid #ddd;border-radius:8px;overflow:hidden;'>
          <embed src="{path_to_show}" type="application/pdf" width="100%" height="{height}px" />
        </div>
        """
        html(html_code, height=height + 10)
    except Exception as e:
        st.error(f"PDF embed error (file URL): {e}")


def extract_sections_from_docx(docx_path: Optional[str] = None, docx_bytes: Optional[bytes] = None) -> Dict[str, str]:
    """Parse the Word doc (from path or bytes) and extract text under six top-level sections (fast view)."""
    global DocxDocument
    if DocxDocument is None:
        try:
            DocxDocument = importlib.import_module("docx").Document
        except Exception:
            return {s: "*(python-docx not available — install with `pip install python-docx`)*" for s in SECTIONS}

    out = {s: [] for s in SECTIONS}
    try:
        if docx_bytes is not None:
            doc = DocxDocument(io.BytesIO(docx_bytes))
        elif docx_path and os.path.exists(docx_path):
            doc = DocxDocument(docx_path)
        else:
            return {s: f"*(DOCX not found at {docx_path or '[uploaded bytes missing]'})*" for s in SECTIONS}
    except Exception as e:
        return {s: f"*(Failed to open DOCX: {e})*" for s in SECTIONS}

    current_section: Optional[str] = None
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        # 1) Heading-based detection
        if "HEADING" in style_name.upper():
            matched = False
            for sec, pattern in PHASE_PATTERNS.items():
                if pattern.search(text):
                    current_section = sec
                    out[sec].append(f"### {text}")
                    matched = True
                    break
            if matched:
                continue
        # 2) Pseudo-heading
        for sec, pattern in PHASE_PATTERNS.items():
            if pattern.search(text) and (text.isupper() or len(text.split()) <= 6):
                current_section = sec
                out[sec].append(f"### {text}")
                break
        else:
            if current_section:
                out[current_section].append(text)
    return {k: ("\n\n".join(v) if v else "*(No section content found — check headings/keywords in the DOCX)*") for k, v in out.items()}


def docx_to_pdf_bytes(docx_path: Optional[str] = None, docx_bytes: Optional[bytes] = None) -> Optional[bytes]:
    """Best-effort convert DOCX to PDF bytes using LibreOffice/soffice or docx2pdf (Word)."""
    # Write input to temp file
    if docx_path and os.path.exists(docx_path):
        in_path = docx_path
        cleanup_in = False
    elif docx_bytes is not None:
        tf_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tf_in.write(docx_bytes)
        tf_in.close()
        in_path = tf_in.name
        cleanup_in = True
    else:
        return None

    out_dir = tempfile.mkdtemp()
    out_pdf = os.path.join(out_dir, "out.pdf")

    # 1) Try soffice/libreoffice
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            import subprocess
            cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, in_path]
            subprocess.check_call(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            # LibreOffice names the output after input filename
            guess = os.path.join(out_dir, os.path.basename(in_path).rsplit(".",1)[0] + ".pdf")
            if os.path.exists(guess):
                out_pdf = guess
            with open(out_pdf, "rb") as f:
                data = f.read()
            if cleanup_in:
                try: os.unlink(in_path)
                except Exception: pass
            return data
        except Exception:
            pass

    # 2) Try docx2pdf (Word on Windows/macOS)
    try:
        if ensure_dependency("docx2pdf", "docx2pdf"):
            import docx2pdf
            tmp_pdf = os.path.join(out_dir, "converted.pdf")
            docx2pdf.convert(in_path, tmp_pdf)
            if os.path.exists(tmp_pdf):
                with open(tmp_pdf, "rb") as f:
                    data = f.read()
                if cleanup_in:
                    try: os.unlink(in_path)
                    except Exception: pass
                return data
    except Exception:
        pass

    if cleanup_in:
        try: os.unlink(in_path)
        except Exception: pass
    # Failed
    return None


def render_docx_as_images(docx_path: Optional[str] = None, docx_bytes: Optional[bytes] = None, zoom: float = 2.0) -> List[bytes]:
    """Return list of PNG page images for the DOCX (via DOCX->PDF->images)."""
    pdf_bytes = docx_to_pdf_bytes(docx_path=docx_path, docx_bytes=docx_bytes)
    if not pdf_bytes:
        return []
    return pdf_to_images(pdf_bytes=pdf_bytes, zoom=zoom)


def docx_to_html(docx_path: Optional[str] = None, docx_bytes: Optional[bytes] = None) -> Optional[str]:
    """Convert DOCX to HTML using Mammoth (best-effort)."""
    global mammoth
    if mammoth is None:
        try:
            mammoth = importlib.import_module("mammoth")
        except Exception:
            return None
    try:
        if docx_bytes is not None:
            result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
        elif docx_path and os.path.exists(docx_path):
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(f)
        else:
            return None
        html_body = result.value  # The generated HTML
        # wrap with minimal CSS for readability
        css = """
        <style>
          body { font-family: ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Arial; line-height: 1.4; }
          h1,h2,h3 { margin: 0.6rem 0; }
          p { margin: 0.4rem 0; }
          table { border-collapse: collapse; width: 100%; }
          th, td { border: 1px solid #ddd; padding: 6px; vertical-align: top; }
        </style>
        """
        return f"<!doctype html><html><head><meta charset='utf-8'>{css}</head><body>{html_body}</body></html>"
    except Exception:
        return None


def show_sidebar_navigation():
    st.sidebar.title("Workshops Navigator")
    st.sidebar.caption("Click a phase to load content from the guidebook.")
    for sec in SECTIONS:
        if st.sidebar.button(sec, use_container_width=True, key=f"nav_{sec}"):
            st.session_state.current_section = sec
    st.sidebar.markdown("---")
    if st.sidebar.button("Show Dashboard PDF", use_container_width=True, key="show_pdf"):
        st.session_state.current_section = None


def diagnostics(pdf_path: Optional[str], docx_path: Optional[str], pdf_bytes: Optional[bytes], docx_bytes: Optional[bytes]):
    st.sidebar.markdown("### Diagnostics")
    def has_mod(name): 
        try: importlib.import_module(name); return True
        except Exception: return False
    st.sidebar.write(f"**python-docx available:** {'✅' if has_mod('docx') else '❌'}")
    st.sidebar.write(f"**PyMuPDF (fitz) available:** {'✅' if has_mod('fitz') else '❌'}")
    st.sidebar.write(f"**Mammoth available:** {'✅' if has_mod('mammoth') else '❌'}")
    st.sidebar.write(f"**PDF path found:** {'✅' if (pdf_path and os.path.exists(pdf_path)) else '—'}")
    st.sidebar.write(f"**PDF uploaded:** {'✅' if pdf_bytes else '—'}")
    st.sidebar.write(f"**DOCX path found:** {'✅' if (docx_path and os.path.exists(docx_path)) else '—'}")
    st.sidebar.write(f"**DOCX uploaded:** {'✅' if docx_bytes else '—'}")


def show_pdf_view(pdf_path: Optional[str], pdf_bytes: Optional[bytes], mode: str, zoom: float, viewer_height: int):
    st.subheader("Facilitator’s Dashboard")
    if pdf_bytes is not None:
        st.caption("Uploaded PDF detected — rendering as images for consistent display.")
        images = pdf_to_images(pdf_bytes=pdf_bytes, zoom=zoom)
        if images:
            for i, img_bytes in enumerate(images, 1):
                st.image(img_bytes, caption=f"Dashboard Page {i}", use_container_width=True)
        else:
            st.error("Could not render images.")
        st.download_button("Download uploaded PDF", data=pdf_bytes, file_name="uploaded.pdf", mime="application/pdf")
        return

    st.caption("No uploaded PDF — using configured path and selected display mode.")
    if mode == "Images only":
        images = pdf_to_images(pdf_path=pdf_path, zoom=zoom)
        if images:
            for i, img_bytes in enumerate(images, 1):
                st.image(img_bytes, caption=f"Dashboard Page {i}", use_container_width=True)
        else:
            st.warning("Could not render images (PyMuPDF missing or invalid PDF).")
    elif mode == "Data-URI only":
        embed_pdf_as_data_uri(pdf_path=pdf_path, height=viewer_height)
    elif mode == "File URL only":
        embed_pdf_as_file_url(pdf_path=pdf_path, height=viewer_height)
    else:  # Auto
        images = pdf_to_images(pdf_path=pdf_path, zoom=zoom)
        if images:
            for i, img_bytes in enumerate(images, 1):
                st.image(img_bytes, caption=f"Dashboard Page {i}", use_container_width=True)
        else:
            try:
                embed_pdf_as_data_uri(pdf_path=pdf_path, height=viewer_height)
            except Exception:
                embed_pdf_as_file_url(pdf_path=pdf_path, height=viewer_height)


def main():
    st.set_page_config(page_title="JIBC Workshops Navigator", page_icon="🗂️", layout="wide", initial_sidebar_state="expanded")

    # One-time dependency bootstrap (best effort)
    if not st.session_state.get("deps_bootstrapped", False):
        ensure_dependency("python-docx", "docx")
        ensure_dependency("pymupdf", "fitz")
        ensure_dependency("mammoth", "mammoth")
        st.session_state["deps_bootstrapped"] = True
        st.rerun()

    if "current_section" not in st.session_state:
        st.session_state.current_section = None

    st.title("JIBC Workshops – Phased Navigator")
    st.write("View the phased dashboard PDF and jump to content sections from the full facilitation guidebook.")

    # Files expander: allow paths or uploads; prefer uploads when provided
    with st.sidebar.expander("Files", expanded=False):
        pdf_path = st.text_input("PDF path", value=DEFAULT_PDF_PATH)
        docx_path = st.text_input("DOCX path", value=DEFAULT_DOCX_PATH)
        uploaded_pdf = st.file_uploader("Or upload a PDF", type=["pdf"], key="pdf_up")
        uploaded_docx = st.file_uploader("Or upload a DOCX", type=["docx"], key="docx_up")

    pdf_bytes = uploaded_pdf.read() if uploaded_pdf is not None else None
    docx_bytes = uploaded_docx.read() if uploaded_docx is not None else None

    # Rendering controls (PDF)
    mode = st.sidebar.radio(
        "PDF display mode (dashboard)",
        options=["Auto (Images → Data-URI → File URL)", "Images only", "Data-URI only", "File URL only"],
    )
    zoom = st.sidebar.slider("Image zoom (PDF & DOCX page images)", min_value=1.0, max_value=3.0, value=2.0, step=0.25)
    viewer_height = st.sidebar.slider("Viewer height (px) for Data-URI/File URL", min_value=500, max_value=1400, value=860, step=20)

    # Rendering controls (DOCX/Guidebook)
    st.sidebar.markdown("---")
    guide_mode = st.sidebar.radio(
        "Guidebook display mode (Word)",
        options=["Extracted text (fast)", "Page images (best fidelity)", "HTML (Mammoth)"],
        index=0,
        help="For full formatting (shapes/boxes), choose Page images (requires LibreOffice/soffice or Word via docx2pdf)."
    )

    diagnostics(pdf_path, docx_path, pdf_bytes, docx_bytes)
    show_sidebar_navigation()

    # Always extract sections for the fast view & nav
    sections_map = extract_sections_from_docx(docx_path if docx_bytes is None else None, docx_bytes=docx_bytes)

    if st.session_state.current_section is None:
        # Show dashboard PDF view
        show_pdf_view(pdf_path if pdf_bytes is None else None, pdf_bytes=pdf_bytes, mode=mode, zoom=zoom, viewer_height=viewer_height)
    else:
        sec = st.session_state.current_section
        st.subheader(f"{sec} — Facilitation Guidebook")
        if guide_mode == "Extracted text (fast)":
            content = sections_map.get(sec, "*(No content found for this section.)*")
            st.markdown(content)
            st.info("Switch to 'Page images' or 'HTML' in the sidebar to see formatted output.")
        elif guide_mode == "HTML (Mammoth)":
            html_doc = docx_to_html(docx_path if docx_bytes is None else None, docx_bytes=docx_bytes)
            if html_doc:
                html(html_doc, height=viewer_height, scrolling=True)
                st.caption("Rendered with Mammoth — boxes/shapes may be flattened.")
            else:
                st.error("Could not convert DOCX to HTML. Ensure 'mammoth' is installed.")
        else:  # Page images (best fidelity)
            st.caption("Rendering full guidebook pages as images for highest fidelity.")
            # We render the WHOLE guidebook with formatting; (not section-filtered)
            imgs = render_docx_as_images(docx_path if docx_bytes is None else None, docx_bytes=docx_bytes, zoom=zoom)
            if imgs:
                for i, im in enumerate(imgs, 1):
                    st.image(im, caption=f"Guidebook Page {i}", use_container_width=True)
            else:
                st.error("Page image rendering failed. Install LibreOffice (soffice in PATH) or run with Word + docx2pdf available.")
                st.info("Tip: Install LibreOffice and restart the app, or switch to 'HTML (Mammoth)' mode.")

    st.sidebar.markdown("---")
    st.sidebar.caption("For perfect fidelity, install LibreOffice locally:\nhttps://www.libreoffice.org/download/download/")

if __name__ == "__main__":
    main()
